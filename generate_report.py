import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

# Import content generators (we'll implement these as files in the same dir)
try:
    import part_front
    import part_ch
except ImportError as e:
    print(f"Error importing parts: {e}")
    sys.exit(1)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def doc_add_heading(doc, text, level, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_paragraph()
    h.alignment = align
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
    h.paragraph_format.space_after = Pt(12)
    h.paragraph_format.space_before = Pt(12)
    return h

def doc_add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    return p

doc = Document()

# Page setup: A4
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    
    # Header format
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "Road Pothole and Damage Detection System Using Deep Learning"
    header_p.style.font.name = 'Times New Roman'
    header_p.style.font.size = Pt(10)
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Global Font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ================= FRONT MATTER =================
fm = part_front.get_front_matter()

# Page 1: Title
p = doc_add_paragraph(doc, fm["title"]["college"] + "\n" + fm["title"]["dept"], WD_ALIGN_PARAGRAPH.CENTER, bold=True)
p.runs[0].font.size = Pt(16)
doc.add_paragraph()
p = doc_add_paragraph(doc, "Mini Project Report\non", WD_ALIGN_PARAGRAPH.CENTER)
p.runs[0].font.size = Pt(14)
doc.add_paragraph()
p = doc_add_paragraph(doc, '"' + fm["title"]["title"] + '"', WD_ALIGN_PARAGRAPH.CENTER, bold=True)
p.runs[0].font.size = Pt(16)
doc.add_paragraph()
doc_add_paragraph(doc, "Submitted in partial fulfillment of the requirements\nfor the degree of\n" + fm["title"]["degree"], WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc_add_paragraph(doc, "Submitted by:\n" + fm["title"]["student"] + "\n\nUnder the Guidance of:\n" + fm["title"]["guide"], WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
doc_add_paragraph(doc, "[College Logo placeholder]", WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc_add_paragraph(doc, fm["title"]["college"] + "\n" + fm["title"]["city"] + "\n[Year: " + fm["title"]["year"] + "]", WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# Page 2: Certificate
doc_add_heading(doc, "CERTIFICATE", 1)
doc.add_paragraph()
doc_add_paragraph(doc, fm["certificate"], WD_ALIGN_PARAGRAPH.JUSTIFY)
doc.add_paragraph()
doc.add_paragraph()
doc_add_paragraph(doc, "Guide: _____________\t\tHead of Department: _____________\nDate:  _____________")
doc.add_page_break()

# Page 3: Declaration
doc_add_heading(doc, "DECLARATION", 1)
doc.add_paragraph()
doc_add_paragraph(doc, fm["declaration"], WD_ALIGN_PARAGRAPH.JUSTIFY)
doc.add_paragraph()
doc.add_paragraph()
doc_add_paragraph(doc, "Neeraj Negi\nDate:\nPlace:")
doc.add_page_break()

# Page 4: Acknowledgement
doc_add_heading(doc, "ACKNOWLEDGEMENT", 1)
doc_add_paragraph(doc, fm["acknowledgement"], WD_ALIGN_PARAGRAPH.JUSTIFY)
doc.add_page_break()

# Page 5: Abstract
doc_add_heading(doc, "ABSTRACT", 1)
doc_add_paragraph(doc, fm["abstract"], WD_ALIGN_PARAGRAPH.JUSTIFY)
doc.add_page_break()

# Page 6: TOC
doc_add_heading(doc, "TABLE OF CONTENTS", 1)
doc_add_paragraph(doc, "Chapter 1: INTRODUCTION .............................................................. 1", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "Chapter 2: LITERATURE REVIEW ......................................................... 5", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "Chapter 3: SYSTEM DESIGN AND METHODOLOGY ............................................. 9", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "Chapter 4: IMPLEMENTATION ............................................................ 15", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "Chapter 5: RESULTS AND ANALYSIS ...................................................... 20", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "Chapter 6: CONCLUSION AND FUTURE WORK ................................................ 24", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "REFERENCES ........................................................................... 26", WD_ALIGN_PARAGRAPH.LEFT)
doc_add_paragraph(doc, "APPENDIX ............................................................................. 27", WD_ALIGN_PARAGRAPH.LEFT)
doc.add_page_break()

# Page 7: LOF
doc_add_heading(doc, "LIST OF FIGURES", 1)
for fig in fm["lists"]["figures"]:
    doc_add_paragraph(doc, fig)
doc.add_page_break()

# Page 8: LOT
doc_add_heading(doc, "LIST OF TABLES", 1)
for tab in fm["lists"]["tables"]:
    doc_add_paragraph(doc, tab)
doc.add_page_break()

# ================= MAIN CONTENT =================
# We add a new section here so we can potentially re-start page numbers (or just add them to the footer)
section = doc.add_section()
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer_p.add_run())

chapters = part_ch.get_chapters()
for ch in chapters:
    # Add Chapter Heading
    doc_add_heading(doc, ch['title'], 1)
    
    # Process blocks
    for block in ch['blocks']:
        if block['type'] == 'heading2':
            doc_add_heading(doc, block['text'], 2)
        elif block['type'] == 'heading3':
            doc_add_heading(doc, block['text'], 3)
        elif block['type'] == 'paragraph':
            doc_add_paragraph(doc, block['text'], WD_ALIGN_PARAGRAPH.JUSTIFY)
        elif block['type'] == 'table':
            # rudimentary table addition
            rows = block['data']
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for idx, r in enumerate(rows):
                rt = table.rows[idx]
                for cidx, cval in enumerate(r):
                    rt.cells[cidx].text = str(cval)
            doc.add_paragraph() # spacing
            
    doc.add_page_break()

# Save document
doc.save("Road_Pothole_Detection_Project_Report.docx")
print("Report generated successfully.")
